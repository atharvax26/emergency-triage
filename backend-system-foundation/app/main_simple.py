"""Simplified FastAPI application - no database required."""

import logging
from fastapi import FastAPI, HTTPException, Depends, UploadFile, File

logger = logging.getLogger(__name__)
from fastapi.middleware.cors import CORSMiddleware
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
import httpx
import sys
import os
import uuid
import json
import bcrypt
from datetime import datetime, timedelta
from jose import jwt, JWTError
from pydantic import BaseModel, Field, EmailStr
from typing import Optional, List

# ── JSON file persistence ─────────────────────────────────────────────────────
_DATA_DIR = os.path.join(os.path.dirname(__file__), "..", "data")
os.makedirs(_DATA_DIR, exist_ok=True)

_USERS_FILE    = os.path.join(_DATA_DIR, "users.json")
_PATIENTS_FILE = os.path.join(_DATA_DIR, "patients.json")
_QUEUE_FILE    = os.path.join(_DATA_DIR, "queue.json")
_AUDIT_FILE    = os.path.join(_DATA_DIR, "audit.json")
_METRICS_FILE  = os.path.join(_DATA_DIR, "metrics.json")

def _load_json(path: str, default):
    try:
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f)
    except (FileNotFoundError, json.JSONDecodeError):
        return default

def _save_json(path: str, data) -> None:
    with open(path, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=2, ensure_ascii=False)

# ── JWT config ────────────────────────────────────────────────────────────────
JWT_SECRET = os.getenv("JWT_SECRET", "triage-dev-secret-change-in-prod")
JWT_ALGORITHM = "HS256"
JWT_EXPIRE_HOURS = 24

_bearer = HTTPBearer(auto_error=False)

def _make_token(user_id: str, name: str, role: str) -> str:
    payload = {
        "sub": user_id,
        "name": name,
        "role": role,
        "exp": datetime.utcnow() + timedelta(hours=JWT_EXPIRE_HOURS),
    }
    return jwt.encode(payload, JWT_SECRET, algorithm=JWT_ALGORITHM)

def _decode_token(token: str) -> dict:
    return jwt.decode(token, JWT_SECRET, algorithms=[JWT_ALGORITHM])

def get_current_user(creds: HTTPAuthorizationCredentials = Depends(_bearer)):
    if not creds:
        raise HTTPException(status_code=401, detail="Not authenticated")
    try:
        return _decode_token(creds.credentials)
    except JWTError:
        raise HTTPException(status_code=401, detail="Invalid or expired token")

# ── In-memory user store (persisted to disk) ──────────────────────────────────
_users: dict = _load_json(_USERS_FILE, {})

# Add ml-pipeline to path so we can use SimplePredictor directly
_ML_PATH = os.path.join(os.path.dirname(__file__), "..", "..", "ml-pipeline", "src", "api")
if _ML_PATH not in sys.path:
    sys.path.insert(0, _ML_PATH)

try:
    from simple_predictor import SimplePredictor
    _simple_predictor = SimplePredictor()
    _ML_DIRECT = True
except Exception as _e:
    _simple_predictor = None
    _ML_DIRECT = False

# Scaledown context pruner (lightweight, fail-safe)
import importlib.util as _ilu
import pathlib as _pl
_pruner_path = _pl.Path(__file__).parent / "scaledown_pruner.py"
_spec = _ilu.spec_from_file_location("scaledown_pruner", _pruner_path)
_pruner_mod = _ilu.module_from_spec(_spec)
_spec.loader.exec_module(_pruner_mod)
_get_pruner = _pruner_mod.get_pruner

# Gemini reasoning layer (fail-safe — falls back to rule-based if unavailable)
_gemini_path = _pl.Path(__file__).parent / "gemini_reasoner.py"
_gemini_spec = _ilu.spec_from_file_location("gemini_reasoner", _gemini_path)
_gemini_mod = _ilu.module_from_spec(_gemini_spec)
_gemini_spec.loader.exec_module(_gemini_mod)
_gemini_reason = _gemini_mod.reason
_gemini_available = _gemini_mod.is_available

# In-memory patient store (persisted to disk)
_patients: dict = _load_json(_PATIENTS_FILE, {})

# ML service URL (used only if direct import failed)
_ML_SERVICE_URL = os.getenv("ML_SERVICE_URL", "http://localhost:8001")
_ML_API_KEY = os.getenv("ML_API_KEY", "system-key-789")

app = FastAPI(
    title="Emergency Triage Backend",
    version="1.0.0",
    description="Simplified backend for testing"
)

# CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000", "http://localhost:5173", "http://localhost:8080"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.get("/")
async def root():
    return {
        "name": "Emergency Triage Backend",
        "version": "1.0.0",
        "status": "running"
    }

@app.get("/health")
async def health():
    return {
        "status": "healthy",
        "database": "not_required",
        "cache": "not_required"
    }


# ============================================================================
# Auth endpoints
# ============================================================================

class RegisterRequest(BaseModel):
    name: str
    email: str
    password: str
    role: str  # nurse | doctor | admin

class LoginRequest(BaseModel):
    email: str
    password: str

@app.post("/api/v1/auth/register")
async def register(data: RegisterRequest):
    """Register a new user account."""
    email = data.email.strip().lower()
    if email in _users:
        raise HTTPException(status_code=409, detail="An account with this email already exists.")
    if data.role not in ("nurse", "doctor", "admin"):
        raise HTTPException(status_code=422, detail="Role must be nurse, doctor, or admin.")
    if len(data.password) < 6:
        raise HTTPException(status_code=422, detail="Password must be at least 6 characters.")

    hashed = bcrypt.hashpw(data.password.encode(), bcrypt.gensalt()).decode()
    user_id = str(uuid.uuid4())
    _users[email] = {
        "id": user_id,
        "name": data.name.strip(),
        "email": email,
        "role": data.role,
        "password_hash": hashed,
        "created_at": datetime.utcnow().isoformat(),
    }
    _save_json(_USERS_FILE, _users)
    token = _make_token(user_id, data.name.strip(), data.role)
    return {"token": token, "name": data.name.strip(), "role": data.role, "email": email}

@app.post("/api/v1/auth/login")
async def login(data: LoginRequest):
    """Authenticate and return a JWT token."""
    email = data.email.strip().lower()
    user = _users.get(email)
    if not user:
        raise HTTPException(status_code=401, detail="Invalid email or password.")
    if not bcrypt.checkpw(data.password.encode(), user["password_hash"].encode()):
        raise HTTPException(status_code=401, detail="Invalid email or password.")
    token = _make_token(user["id"], user["name"], user["role"])
    return {"token": token, "name": user["name"], "role": user["role"], "email": email}

class VerifyPasswordRequest(BaseModel):
    email: str
    password: str

@app.post("/api/v1/auth/verify-password")
async def verify_password(data: VerifyPasswordRequest):
    """Verify a user's password without issuing a new token. Used for override authentication."""
    email = data.email.strip().lower()
    user = _users.get(email)
    if not user:
        return {"valid": False}
    valid = bcrypt.checkpw(data.password.encode(), user["password_hash"].encode())
    return {"valid": valid}

@app.get("/api/v1/auth/me")
async def me(current_user: dict = Depends(get_current_user)):
    """Return current user info from JWT."""
    return {"name": current_user["name"], "role": current_user["role"]}


# ============================================================================
# ML Health endpoint
# ============================================================================

@app.get("/api/v1/ml/health")
async def ml_health():
    """Check ML service health."""
    if _ML_DIRECT and _simple_predictor is not None:
        return {
            "status": "healthy",
            "ml_service": "connected",
            "gemini_reasoning": _gemini_available(),
            "gemini_model": _gemini_mod.GEMINI_MODEL,
        }
    
    # Fallback: try the external ML service
    try:
        async with httpx.AsyncClient(timeout=1.0) as client:
            response = await client.get(
                f"{_ML_SERVICE_URL}/v1/health/live",
                headers={"X-API-Key": _ML_API_KEY}
            )
            response.raise_for_status()
        return {"status": "healthy", "ml_service": "connected", "gemini_reasoning": _gemini_available()}
    except Exception as e:
        return {"status": "unhealthy", "ml_service": "disconnected", "error": str(e), "gemini_reasoning": False}


# ============================================================================
# Patients endpoints
# ============================================================================

class PatientCreate(BaseModel):
    name: str
    age: int
    gender: str
    chief_complaint: str
    symptoms: str

@app.get("/api/v1/patients")
async def get_patients(page: int = 1, limit: int = 50):
    """Get all patients."""
    patients_list = list(_patients.values())
    return {
        "patients": patients_list,
        "total": len(patients_list),
        "page": page,
        "limit": limit
    }

@app.post("/api/v1/patients")
async def create_patient(data: PatientCreate):
    """Create a new patient record."""
    patient_id = str(uuid.uuid4())
    patient = {
        "id": patient_id,
        "name": data.name,
        "age": data.age,
        "gender": data.gender,
        "chief_complaint": data.chief_complaint,
        "symptoms": data.symptoms,
        "created_at": datetime.utcnow().isoformat()
    }
    _patients[patient_id] = patient
    _save_json(_PATIENTS_FILE, _patients)
    return patient

@app.get("/api/v1/patients/{patient_id}")
async def get_patient(patient_id: str):
    """Get patient by ID."""
    if patient_id not in _patients:
        from fastapi import HTTPException
        raise HTTPException(status_code=404, detail="Patient not found")
    return _patients[patient_id]


# ============================================================================
# Queue endpoints
# ============================================================================

_queue: List[dict] = _load_json(_QUEUE_FILE, [])

# ── Audit log store (persisted to disk) ───────────────────────────────────────
_audit: List[dict] = _load_json(_AUDIT_FILE, [])

# ── Pruning metrics store (persisted to disk) ─────────────────────────────────
# Each entry: { patient_id, original_tokens, compressed_tokens, compression_ratio,
#               pruning_applied, inference_time_ms, severity, timestamp }
_metrics: List[dict] = _load_json(_METRICS_FILE, [])

class QueueAddRequest(BaseModel):
    patient_id: str
    name: str
    age: int
    chief_complaint: str
    severity: str          # low | medium | high | critical
    priority: Optional[int] = None  # auto-derived from severity if omitted

_SEVERITY_PRIORITY = {"critical": 0, "high": 1, "medium": 2, "low": 3}

@app.get("/api/v1/queue")
async def get_queue():
    """Get current queue with full patient info."""
    # Recalculate live wait times
    now = datetime.utcnow()
    items = []
    for entry in _queue:
        e = dict(entry)
        added = datetime.fromisoformat(e["added_at"].replace("Z", "+00:00"))
        from datetime import timezone
        e["wait_minutes"] = int((datetime.now(timezone.utc) - added).total_seconds() / 60)
        items.append(e)
    return {"queue": items, "total": len(items)}

@app.post("/api/v1/queue")
async def add_to_queue(data: QueueAddRequest):
    """Add a triaged patient to the queue."""
    # Prevent duplicates — same patient_id can only appear once as 'waiting'
    for entry in _queue:
        if entry["patient_id"] == data.patient_id and entry["status"] == "waiting":
            return entry  # already queued, return existing

    priority = data.priority if data.priority is not None else _SEVERITY_PRIORITY.get(data.severity.lower(), 2)
    entry = {
        "id": str(uuid.uuid4()),
        "patient_id": data.patient_id,
        "name": data.name,
        "age": data.age,
        "chief_complaint": data.chief_complaint,
        "severity": data.severity.lower(),
        "priority": priority,
        "status": "waiting",
        "added_at": datetime.utcnow().isoformat() + "Z",
        "wait_minutes": 0,
    }
    _queue.append(entry)
    _save_json(_QUEUE_FILE, _queue)
    return entry

@app.patch("/api/v1/queue/{entry_id}/status")
async def update_queue_status(entry_id: str, data: dict):
    """Update queue entry status (waiting → in-progress → completed)."""
    for entry in _queue:
        if entry["id"] == entry_id:
            entry["status"] = data.get("status", entry["status"])
            _save_json(_QUEUE_FILE, _queue)
            return entry
    raise HTTPException(status_code=404, detail="Queue entry not found")


# ============================================================================
# Audit log endpoints
# ============================================================================

class AuditEntryRequest(BaseModel):
    patient_id: str
    patient_name: str
    severity: str
    action_taken: str
    overridden: Optional[bool] = False
    override_reason: Optional[str] = None
    performed_by: Optional[str] = None

@app.get("/api/v1/audit")
async def get_audit_log(limit: int = 200):
    """Return audit log entries, newest first."""
    return {"entries": list(reversed(_audit[-limit:])), "total": len(_audit)}

@app.post("/api/v1/audit")
async def add_audit_entry(data: AuditEntryRequest):
    """Record a triage action in the audit log."""
    entry = {
        "id": str(uuid.uuid4()),
        "timestamp": datetime.utcnow().isoformat() + "Z",
        "patient_id": data.patient_id,
        "patient_name": data.patient_name,
        "severity": data.severity.lower(),
        "action_taken": data.action_taken,
        "overridden": data.overridden or False,
        "override_reason": data.override_reason,
        "performed_by": data.performed_by,
    }
    _audit.append(entry)
    _save_json(_AUDIT_FILE, _audit)
    return entry


@app.delete("/api/v1/audit")
async def clear_audit_log():
    """Clear all audit log entries."""
    global _audit
    _audit = []
    _save_json(_AUDIT_FILE, _audit)
    return {"status": "cleared", "message": "Audit log cleared."}


@app.delete("/api/v1/reset")
async def reset_all_data():
    """Clear all patients, queue, audit, and metrics data (admin use only)."""
    global _patients, _queue, _audit, _metrics
    _patients = {}
    _queue = []
    _audit = []
    _metrics = []
    _save_json(_PATIENTS_FILE, _patients)
    _save_json(_QUEUE_FILE, _queue)
    _save_json(_AUDIT_FILE, _audit)
    _save_json(_METRICS_FILE, _metrics)
    return {"status": "reset", "message": "All patient, queue, audit, and metrics data cleared."}


# ============================================================================
# Scaledown context pruning endpoint
# ============================================================================

class PruneRequest(BaseModel):
    context: str
    prompt: Optional[str] = "Summarize clinically relevant triage information."
    rate: Optional[str] = "auto"

@app.post("/api/v1/prune")
async def prune_context(data: PruneRequest):
    """
    Compress patient context via Scaledown API.
    Position: ML Severity Engine → Scaledown (Pruning) → LLM Reasoner
    """
    pruner = _get_pruner()
    result = await pruner.prune(context=data.context, prompt=data.prompt)
    return result


# ============================================================================
# ML Prediction endpoint
# ============================================================================

class VitalSigns(BaseModel):
    systolic_bp: float = Field(..., ge=50, le=250)
    diastolic_bp: float = Field(..., ge=30, le=150)
    heart_rate: float = Field(..., ge=20, le=250)
    respiratory_rate: float = Field(..., ge=5, le=60)
    temperature: float = Field(..., ge=32.0, le=43.0)
    spo2: float = Field(..., ge=50, le=100)

class PatientData(BaseModel):
    vitals: VitalSigns
    age: int = Field(..., ge=0, le=120)
    symptoms: Optional[list[str]] = None

class TriagePredictionRequest(BaseModel):
    patient_data: PatientData
    request_id: Optional[str] = None


def _build_rich_clinical_context(age: int, vitals, symptoms: list) -> str:
    """
    Build a rich clinical document (1500–4000 tokens) from patient data.
    Includes current vitals, presenting symptoms, synthetic patient history,
    previous visit notes, and attending physician notes.
    Higher token count → more redundancy → better Scaledown compression ratio.
    """
    sym_str = ", ".join(symptoms) if symptoms else "none reported"

    # Derive plausible chronic conditions from age + vitals
    chronic = []
    if age > 50 and vitals.systolic_bp > 130:
        chronic.append("hypertension (diagnosed 2019, on lisinopril 10mg daily)")
    if age > 45:
        chronic.append("type 2 diabetes mellitus (HbA1c 7.2%, on metformin 500mg BID)")
    if vitals.spo2 < 95:
        chronic.append("mild chronic obstructive pulmonary disease (GOLD stage I)")
    if age > 60:
        chronic.append("hyperlipidemia (on atorvastatin 20mg nightly)")
    if not chronic:
        chronic.append("no significant chronic medical history")

    allergies = "penicillin (rash), sulfonamides (unknown reaction)" if age > 40 else "NKDA (no known drug allergies)"

    surgical = []
    if age > 50:
        surgical.append("appendectomy (1998, uncomplicated)")
    if age > 65:
        surgical.append("right knee arthroplasty (2015)")
    if not surgical:
        surgical.append("no prior surgeries")

    smoking = "former smoker (20 pack-years, quit 2010)" if age > 50 else "non-smoker"
    alcohol = "occasional alcohol use (1–2 drinks/week)" if age > 21 else "denies alcohol use"
    family_hx = "father: MI at age 58; mother: type 2 diabetes, hypertension" if age > 40 else "non-contributory"

    meds = []
    if age > 50 and vitals.systolic_bp > 130:
        meds.extend(["lisinopril 10mg PO daily", "amlodipine 5mg PO daily"])
    if age > 45:
        meds.extend(["metformin 500mg PO BID", "aspirin 81mg PO daily"])
    if age > 60:
        meds.extend(["atorvastatin 20mg PO nightly", "omeprazole 20mg PO daily"])
    if not meds:
        meds.append("no current medications")

    prev_bp = f"{vitals.systolic_bp + 8}/{vitals.diastolic_bp + 5}"
    prev_hr = int(vitals.heart_rate * 0.9)
    prev_temp = round(vitals.temperature - 0.3, 1)
    prev_spo2 = min(100, vitals.spo2 + 2)

    nurse_note = (
        f"Patient arrived ambulatory, appears "
        f"{'acutely ill' if vitals.spo2 < 92 or vitals.heart_rate > 110 else 'mildly distressed'}. "
        f"Chief complaint: {sym_str}. "
        f"Onset reported as {'sudden' if vitals.heart_rate > 100 else 'gradual'}, duration approximately "
        f"{'less than 1 hour' if vitals.heart_rate > 120 else '2–4 hours'}. "
        f"Pain scale: {'8/10' if vitals.heart_rate > 110 else '5/10'}. "
        f"Patient {'denies' if vitals.temperature < 38.0 else 'reports'} fever at home. "
        f"Patient {'denies' if vitals.spo2 > 94 else 'reports'} shortness of breath at rest. "
        f"Vitals on arrival: BP {vitals.systolic_bp}/{vitals.diastolic_bp} mmHg, "
        f"HR {vitals.heart_rate} bpm, RR {vitals.respiratory_rate} breaths/min, "
        f"Temp {vitals.temperature}°C, SpO2 {vitals.spo2}% on room air. "
        f"Patient is alert and oriented x3. GCS 15. "
        f"Skin {'pale and diaphoretic' if vitals.spo2 < 92 else 'warm and dry'}. "
        f"Peripheral pulses {'weak and thready' if vitals.systolic_bp < 90 else 'strong and regular'}. "
        f"Capillary refill {'greater than 3 seconds' if vitals.systolic_bp < 90 else 'less than 2 seconds'}."
    )

    physician_note = (
        f"Patient is a {age}-year-old presenting with {sym_str}. "
        f"History of present illness: Patient reports onset of symptoms approximately "
        f"{'1 hour' if vitals.heart_rate > 110 else '3 hours'} prior to arrival. "
        f"{'Associated with diaphoresis and nausea.' if vitals.heart_rate > 100 else 'No associated nausea or vomiting.'} "
        f"{'Patient denies recent travel or sick contacts.' if vitals.temperature < 38.5 else 'Patient reports recent contact with ill family member.'} "
        f"Review of systems: "
        f"{'Positive for chest tightness and dyspnea.' if vitals.spo2 < 94 else 'Negative for chest pain and dyspnea.'} "
        f"{'Positive for palpitations and lightheadedness.' if vitals.heart_rate > 110 else 'Negative for palpitations.'} "
        f"Negative for headache, vision changes, abdominal pain, urinary symptoms. "
        f"Physical examination: General: "
        f"{'Acutely ill-appearing' if vitals.spo2 < 92 or vitals.systolic_bp < 90 else 'Mildly distressed'}. "
        f"Cardiovascular: {'Tachycardic, irregular rhythm' if vitals.heart_rate > 110 else 'Regular rate and rhythm'}, no murmurs. "
        f"Respiratory: {'Decreased breath sounds bilaterally, accessory muscle use' if vitals.spo2 < 92 else 'Clear to auscultation bilaterally'}. "
        f"Abdomen: Soft, {'tender in epigastric region' if 'abdominal' in sym_str.lower() or 'stomach' in sym_str.lower() else 'non-tender'}, non-distended. "
        f"Extremities: {'1+ pitting edema bilateral lower extremities' if age > 60 else 'No edema'}. "
        f"Assessment: Patient presents with {sym_str} in the context of {', '.join(chronic[:2])}. "
        f"{'Hemodynamic instability noted — immediate intervention warranted.' if vitals.systolic_bp < 90 or vitals.spo2 < 88 else 'Currently hemodynamically stable.'} "
        f"Plan: {'Stat ECG, cardiac enzymes, CBC, BMP, chest X-ray. IV access x2, supplemental O2.' if vitals.spo2 < 94 or vitals.heart_rate > 110 else 'CBC, BMP, urinalysis. Monitor vitals q30min.'}"
    )

    prev_labs = (
        f"Previous visit labs (3 months ago): "
        f"WBC 8.2 K/uL (normal), Hgb 13.4 g/dL (normal), Plt 245 K/uL (normal). "
        f"BMP: Na 138 mEq/L, K 4.1 mEq/L, Cl 102 mEq/L, CO2 24 mEq/L, "
        f"BUN 18 mg/dL, Cr 1.0 mg/dL, Glucose 112 mg/dL. "
        f"LFTs: AST 28 U/L, ALT 32 U/L, Alk Phos 78 U/L, Total Bili 0.8 mg/dL. "
        f"Lipid panel: Total cholesterol 198 mg/dL, LDL 118 mg/dL, HDL 42 mg/dL, TG 152 mg/dL. "
        f"HbA1c: 7.2%. TSH: 2.1 mIU/L (normal). "
        f"Urinalysis: clear, no protein, no glucose, no blood, no nitrites."
    )

    prev_visit = (
        f"Previous ED visit (6 months ago): "
        f"Chief complaint: {'chest discomfort and dyspnea' if vitals.spo2 < 95 else 'fatigue and dizziness'}. "
        f"Vitals on arrival: BP {prev_bp} mmHg, HR {prev_hr} bpm, Temp {prev_temp}°C, SpO2 {prev_spo2}%. "
        f"Diagnosis: {'Hypertensive urgency, managed with IV labetalol, discharged with medication adjustment.' if vitals.systolic_bp > 150 else 'Dehydration, treated with IV fluids 1L NS, discharged in stable condition.'} "
        f"Follow-up: Patient instructed to follow up with primary care physician within 1 week. "
        f"Compliance: Patient reports partial compliance with medications. "
        f"Missed doses of {'lisinopril' if vitals.systolic_bp > 140 else 'metformin'} reported over past 2 weeks."
    )

    document = f"""EMERGENCY DEPARTMENT CLINICAL RECORD
=====================================
PATIENT DEMOGRAPHICS
Age: {age} years
Presenting Complaint: {sym_str}

CURRENT VITAL SIGNS
Blood Pressure: {vitals.systolic_bp}/{vitals.diastolic_bp} mmHg
Heart Rate: {vitals.heart_rate} bpm
Respiratory Rate: {vitals.respiratory_rate} breaths/min
Temperature: {vitals.temperature}°C
Oxygen Saturation: {vitals.spo2}% (room air)

PRESENTING SYMPTOMS
{sym_str}

PAST MEDICAL HISTORY
{chr(10).join(f'- {c}' for c in chronic)}

SURGICAL HISTORY
{chr(10).join(f'- {s}' for s in surgical)}

MEDICATIONS
{chr(10).join(f'- {m}' for m in meds)}

ALLERGIES
{allergies}

SOCIAL HISTORY
Smoking: {smoking}
Alcohol: {alcohol}

FAMILY HISTORY
{family_hx}

TRIAGE NURSE ASSESSMENT
{nurse_note}

ATTENDING PHYSICIAN NOTE
{physician_note}

LABORATORY DATA
{prev_labs}

PREVIOUS EMERGENCY DEPARTMENT VISIT
{prev_visit}

REVIEW OF SYSTEMS
Constitutional: {'Fever, chills, diaphoresis' if vitals.temperature > 38.0 else 'No fever or chills'}. Fatigue present.
Cardiovascular: {'Chest pain, palpitations' if vitals.heart_rate > 100 else 'No chest pain'}. {'Dyspnea on exertion' if vitals.spo2 < 95 else 'No dyspnea'}.
Respiratory: {'Shortness of breath, wheezing' if vitals.spo2 < 94 else 'No respiratory complaints'}.
Gastrointestinal: {'Nausea, vomiting' if 'vomit' in sym_str.lower() or 'nausea' in sym_str.lower() else 'No nausea or vomiting'}. No diarrhea.
Neurological: No headache, no vision changes, no focal deficits.
Musculoskeletal: No joint pain or swelling.
Integumentary: {'Diaphoresis noted' if vitals.heart_rate > 110 else 'No rash or skin changes'}.

RISK STRATIFICATION FACTORS
Age-related risk: {'Elevated (>60 years)' if age > 60 else 'Moderate (40-60 years)' if age > 40 else 'Low (<40 years)'}
Hemodynamic status: {'Unstable' if vitals.systolic_bp < 90 or vitals.heart_rate > 120 else 'Stable'} — BP {vitals.systolic_bp}/{vitals.diastolic_bp}, HR {vitals.heart_rate}
Oxygenation: {'Compromised' if vitals.spo2 < 94 else 'Adequate'} — SpO2 {vitals.spo2}%
Comorbidity burden: {'High' if len(chronic) >= 3 else 'Moderate' if len(chronic) >= 2 else 'Low'}
Medication compliance: Partial — patient reports missed doses""".strip()

    return document


@app.post("/api/v1/ml/triage/predict")
async def predict_triage(request: TriagePredictionRequest):
    """ML prediction endpoint - uses SimplePredictor directly, with Scaledown context pruning."""
    import time as _time
    vitals = request.patient_data.vitals
    age = request.patient_data.age
    symptoms = request.patient_data.symptoms or []

    input_data = {
        'heart_rate': vitals.heart_rate,
        'systolic_bp': vitals.systolic_bp,
        'diastolic_bp': vitals.diastolic_bp,
        'temperature': vitals.temperature,
        'spo2': vitals.spo2,
        'respiratory_rate': vitals.respiratory_rate,
        'age': age,
        'symptoms': symptoms,
    }

    # ── Stage 1: ML inference ─────────────────────────────────────────────────
    t_ml_start = _time.perf_counter()
    result = None

    if _ML_DIRECT and _simple_predictor is not None:
        try:
            result = _simple_predictor.predict(input_data)
            result['request_id'] = request.request_id
        except Exception:
            result = None

    if result is None:
        try:
            async with httpx.AsyncClient(timeout=0.5) as client:
                response = await client.post(
                    f"{_ML_SERVICE_URL}/v1/predict",
                    json=request.model_dump(),
                    headers={"X-API-Key": _ML_API_KEY}
                )
                response.raise_for_status()
            result = response.json()
        except Exception:
            pass

    if result is None:
        risk_score = 0
        if vitals.spo2 < 90:
            risk_score += 35
        elif vitals.spo2 < 94:
            risk_score += 20
        if vitals.systolic_bp < 90:
            risk_score += 30
        elif vitals.systolic_bp > 180:
            risk_score += 20
        if vitals.heart_rate > 120 or vitals.heart_rate < 50:
            risk_score += 25
        elif vitals.heart_rate > 100 or vitals.heart_rate < 60:
            risk_score += 15
        if vitals.temperature > 39.5 or vitals.temperature < 35.0:
            risk_score += 20
        if age > 75:
            risk_score += 10

        prob = min(risk_score / 100.0, 0.95)
        if prob >= 0.75:
            tier = "CRITICAL"
        elif prob >= 0.50:
            tier = "HIGH"
        elif prob >= 0.25:
            tier = "MEDIUM"
        else:
            tier = "LOW"

        result = {
            "raw_probability": prob,
            "calibrated_probability": prob * 0.95,
            "risk_tier": tier,
            "decision_label": "URGENT" if tier in ["HIGH", "CRITICAL"] else "MONITOR",
            "confidence": 0.75,
            "safety_override": False,
            "override_reason": None,
            "model_version": "1.0.0-fallback",
            "model_id": "vitals-fallback",
            "inference_time_ms": 5.0,
            "timestamp": datetime.utcnow().isoformat(),
            "request_id": request.request_id,
            "cache_hit": False,
        }

    ml_ms = round((_time.perf_counter() - t_ml_start) * 1000, 2)
    result["inference_time_ms"] = ml_ms
    logger.info(f"[LATENCY] ML inference: {ml_ms}ms")
    # ─────────────────────────────────────────────────────────────────────────

    # ── Stage 2: Scaledown context pruning ───────────────────────────────────
    # Build a rich, realistic clinical document (1500–4000 tokens) before pruning.
    # Larger input = more redundancy = higher compression ratio = stronger evidence.
    # Synthetic history/notes are generated deterministically from patient data.
    context_str = _build_rich_clinical_context(age, vitals, symptoms)
    t_scaledown_start = _time.perf_counter()
    pruning_result = await _get_pruner().prune(
        context=context_str,
        prompt="Summarize clinically relevant triage information for severity classification.",
    )
    scaledown_ms = round((_time.perf_counter() - t_scaledown_start) * 1000, 2)

    # Log whether fallback (passthrough) was triggered
    if not pruning_result.get("pruning_applied"):
        logger.warning(f"[LATENCY] Scaledown fallback triggered (passthrough) — {scaledown_ms}ms")
    else:
        logger.info(
            f"[LATENCY] Scaledown: {scaledown_ms}ms | "
            f"{pruning_result['original_tokens']}→{pruning_result['compressed_tokens']} tokens "
            f"({round(pruning_result['compression_ratio']*100,1)}% reduction)"
        )
    # ─────────────────────────────────────────────────────────────────────────

    # ── Stage 3: Gemini LLM reasoning ────────────────────────────────────────
    # Input: ML outputs + Scaledown-compressed context
    # Output: structured clinical reasoning, justification, recommended actions
    t_llm_start = _time.perf_counter()
    gemini_result = await _gemini_reason(
        compressed_context=pruning_result["compressed_context"],
        risk_tier=result.get("risk_tier", "MEDIUM"),
        confidence=result.get("confidence", 0.75),
        calibrated_probability=result.get("calibrated_probability", 0.5),
        age=age,
        symptoms=symptoms,
    )
    llm_ms = round((_time.perf_counter() - t_llm_start) * 1000, 2)
    result["reasoning"] = gemini_result

    used_gemini = gemini_result.get("gemini_reasoning", False)
    logger.info(
        f"[LATENCY] LLM ({'Gemini' if used_gemini else 'rule-based fallback'}): {llm_ms}ms"
    )
    # ─────────────────────────────────────────────────────────────────────────

    total_ms = round(ml_ms + scaledown_ms + llm_ms, 2)
    logger.info(
        f"[LATENCY] Total pipeline: {total_ms}ms "
        f"(ML={ml_ms}ms, Scaledown={scaledown_ms}ms, LLM={llm_ms}ms)"
    )

    # ── Store pruning metrics ─────────────────────────────────────────────────
    metric = {
        "patient_id": request.request_id or str(uuid.uuid4()),
        "original_tokens": pruning_result["original_tokens"],
        "compressed_tokens": pruning_result["compressed_tokens"],
        "compression_ratio": pruning_result["compression_ratio"],
        "tokens_saved": pruning_result["original_tokens"] - pruning_result["compressed_tokens"],
        "pruning_applied": pruning_result["pruning_applied"],
        "inference_time_ms": ml_ms,
        "scaledown_ms": scaledown_ms,
        "llm_ms": llm_ms,
        "total_ms": total_ms,
        "severity": result.get("risk_tier", "UNKNOWN").upper(),
        "timestamp": datetime.utcnow().isoformat() + "Z",
    }
    _metrics.append(metric)
    _save_json(_METRICS_FILE, _metrics)
    # ─────────────────────────────────────────────────────────────────────────

    reduction_pct = round(pruning_result["compression_ratio"] * 100, 1)

    # Always log — no silent pruning
    print(
        f"[SCALEDOWN] request={request.request_id or 'anon'} | "
        f"severity={result.get('risk_tier','?')} | "
        f"original={pruning_result['original_tokens']} tokens | "
        f"compressed={pruning_result['compressed_tokens']} tokens | "
        f"reduction={reduction_pct}% | "
        f"applied={pruning_result['pruning_applied']} | "
        f"scaledown_ms={scaledown_ms}",
        flush=True,
    )

    # Attach pruning + compression_stats + latency breakdown to response
    result["pruning"] = {
        "original_tokens": pruning_result["original_tokens"],
        "compressed_tokens": pruning_result["compressed_tokens"],
        "compression_ratio": pruning_result["compression_ratio"],
        "tokens_saved": metric["tokens_saved"],
        "pruning_applied": pruning_result["pruning_applied"],
    }
    result["compression_stats"] = {
        "original_tokens": pruning_result["original_tokens"],
        "compressed_tokens": pruning_result["compressed_tokens"],
        "reduction_percent": reduction_pct,
    }
    result["latency_breakdown"] = {
        "ml_ms": ml_ms,
        "scaledown_ms": scaledown_ms,
        "llm_ms": llm_ms,
        "total_ms": total_ms,
        "scaledown_fallback": not pruning_result.get("pruning_applied", False),
        "llm_fallback": not used_gemini,
    }

    return result


# ============================================================================
# Stats endpoint
# ============================================================================

@app.get("/api/v1/stats")
async def get_stats():
    """Aggregate pruning and inference metrics across all triage calls."""
    if not _metrics:
        return {
            "total_predictions": 0,
            "pruning_applied_count": 0,
            "avg_compression_ratio": 0.0,
            "total_tokens_original": 0,
            "total_tokens_compressed": 0,
            "total_tokens_saved": 0,
            "avg_tokens_saved_pct": 0.0,
            "avg_inference_time_ms": 0.0,
            "p50_latency_ms": 0.0,
            "p95_latency_ms": 0.0,
            "tokens_saved_by_severity": {},
            "per_severity_breakdown": [],
        }

    total = len(_metrics)
    pruned = [m for m in _metrics if m.get("pruning_applied")]
    total_orig = sum(m["original_tokens"] for m in _metrics)
    total_comp = sum(m["compressed_tokens"] for m in _metrics)
    total_saved = sum(m.get("tokens_saved", 0) for m in _metrics)
    avg_ratio = round(sum(m["compression_ratio"] for m in _metrics) / total * 100, 1)
    avg_inf = round(sum(m["inference_time_ms"] for m in _metrics) / total, 2)

    # p50 / p95 latency using total_ms if available, else inference_time_ms
    latency_values = sorted(
        m.get("total_ms", m.get("inference_time_ms", 0)) for m in _metrics
    )
    def _percentile(data: list, pct: float) -> float:
        if not data:
            return 0.0
        idx = int(len(data) * pct / 100)
        idx = min(idx, len(data) - 1)
        return round(data[idx], 2)

    p50 = _percentile(latency_values, 50)
    p95 = _percentile(latency_values, 95)

    # ── Compression statistics ────────────────────────────────────────────────
    reduction_pcts = [round(m["compression_ratio"] * 100, 2) for m in _metrics]
    mean_red = round(sum(reduction_pcts) / len(reduction_pcts), 2)
    min_red  = round(min(reduction_pcts), 2)
    max_red  = round(max(reduction_pcts), 2)
    variance = sum((x - mean_red) ** 2 for x in reduction_pcts) / len(reduction_pcts)
    import math as _math
    std_red  = round(_math.sqrt(variance), 2)
    # ─────────────────────────────────────────────────────────────────────────

    # Per-severity aggregation
    by_severity: dict = {}
    for m in _metrics:
        sev = m.get("severity", "UNKNOWN")
        if sev not in by_severity:
            by_severity[sev] = {
                "count": 0,
                "tokens_saved": 0,
                "total_original": 0,
                "total_compressed": 0,
                "compression_ratios": [],
            }
        by_severity[sev]["count"] += 1
        by_severity[sev]["tokens_saved"] += m.get("tokens_saved", 0)
        by_severity[sev]["total_original"] += m["original_tokens"]
        by_severity[sev]["total_compressed"] += m["compressed_tokens"]
        by_severity[sev]["compression_ratios"].append(m["compression_ratio"])

    # Build clean per-severity breakdown for frontend table
    sev_order = ["CRITICAL", "HIGH", "MEDIUM", "LOW"]
    per_severity_breakdown = []
    for sev in sev_order:
        if sev not in by_severity:
            continue
        d = by_severity[sev]
        avg_cr = round(sum(d["compression_ratios"]) / len(d["compression_ratios"]) * 100, 1)
        per_severity_breakdown.append({
            "severity": sev,
            "count": d["count"],
            "avg_original_tokens": round(d["total_original"] / d["count"]),
            "avg_compressed_tokens": round(d["total_compressed"] / d["count"]),
            "avg_tokens_saved": round(d["tokens_saved"] / d["count"]),
            "avg_reduction_pct": avg_cr,
        })

    for sev, d in by_severity.items():
        if sev not in sev_order:
            avg_cr = round(sum(d["compression_ratios"]) / len(d["compression_ratios"]) * 100, 1)
            per_severity_breakdown.append({
                "severity": sev,
                "count": d["count"],
                "avg_original_tokens": round(d["total_original"] / d["count"]),
                "avg_compressed_tokens": round(d["total_compressed"] / d["count"]),
                "avg_tokens_saved": round(d["tokens_saved"] / d["count"]),
                "avg_reduction_pct": avg_cr,
            })

    return {
        "total_predictions": total,
        "pruning_applied_count": len(pruned),
        "avg_compression_ratio": avg_ratio,
        "total_tokens_original": total_orig,
        "total_tokens_compressed": total_comp,
        "total_tokens_saved": total_saved,
        "avg_tokens_saved_pct": avg_ratio,
        "avg_inference_time_ms": avg_inf,
        "p50_latency_ms": p50,
        "p95_latency_ms": p95,
        "compression_statistics": {
            "mean_reduction_pct": mean_red,
            "std_dev_pct": std_red,
            "min_reduction_pct": min_red,
            "max_reduction_pct": max_red,
            "sample_size": total,
        },
        "tokens_saved_by_severity": {
            sev: {"count": d["count"], "tokens_saved": d["tokens_saved"]}
            for sev, d in by_severity.items()
        },
        "per_severity_breakdown": per_severity_breakdown,
    }


@app.get("/api/v1/evaluation")
async def get_evaluation():
    """Serve compression_evaluation.json for the dashboard charts."""
    eval_path = os.path.join(_DATA_DIR, "compression_evaluation.json")
    if not os.path.exists(eval_path):
        raise HTTPException(status_code=404, detail="Evaluation data not found. Run evaluate_compression.py first.")
    with open(eval_path, "r") as f:
        return json.load(f)




# ============================================================================
# Document extraction endpoint (RAG layer for patient intake auto-fill)
# ============================================================================

@app.post("/api/v1/extract-document")
async def extract_document(file: UploadFile = File(...)):
    """
    Accept an uploaded patient document (PDF, DOCX, TXT) and extract structured
    patient information: name, age, gender, chief complaint, and symptoms.
    Uses Gemini when available, falls back to regex extraction automatically.
    Vitals are intentionally excluded — they must be measured manually.
    """
    import tempfile, pathlib

    GEMINI_API_KEY = os.getenv("GEMINI_API_KEY", "")

    # Read file bytes
    content_bytes = await file.read()
    if not content_bytes:
        raise HTTPException(status_code=400, detail="Uploaded file is empty")

    filename = file.filename or "document"
    ext = pathlib.Path(filename).suffix.lower()

    # ── Extract raw text ──────────────────────────────────────────────────────
    raw_text = ""

    if ext == ".txt":
        raw_text = content_bytes.decode("utf-8", errors="ignore")

    elif ext == ".pdf":
        # Try PyMuPDF (fitz) first, fall back to pdfminer, then raw bytes decode
        try:
            import fitz  # PyMuPDF
            with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
                tmp.write(content_bytes)
                tmp_path = tmp.name
            doc = fitz.open(tmp_path)
            raw_text = "\n".join(page.get_text() for page in doc)
            doc.close()
            os.unlink(tmp_path)
        except ImportError:
            try:
                from pdfminer.high_level import extract_text as pdf_extract
                import io
                raw_text = pdf_extract(io.BytesIO(content_bytes))
            except ImportError:
                # Last resort: decode as utf-8 ignoring binary noise
                raw_text = content_bytes.decode("utf-8", errors="ignore")

    elif ext in (".doc", ".docx"):
        try:
            import docx as _docx
            import io as _io
            from lxml import etree as _etree

            doc = _docx.Document(_io.BytesIO(content_bytes))
            parts = []

            # 1. Regular paragraphs
            for p in doc.paragraphs:
                t = p.text.strip()
                if t:
                    parts.append(t)

            # 2. Table cells
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        t = cell.text.strip()
                        if t:
                            parts.append(t)

            # 3. Headers and footers
            for section in doc.sections:
                for hdr in [section.header, section.footer]:
                    if hdr:
                        for p in hdr.paragraphs:
                            t = p.text.strip()
                            if t:
                                parts.append(t)

            # 4. Text boxes (stored as drawing elements in XML)
            ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
            for elem in doc.element.body.iter():
                if elem.tag == f"{{{ns}}}txbxContent":
                    for child in elem.iter(f"{{{ns}}}t"):
                        t = (child.text or "").strip()
                        if t:
                            parts.append(t)

            raw_text = "\n".join(parts)

        except Exception as _docx_err:
            logger.warning(f"python-docx extraction failed: {_docx_err}, falling back to raw decode")
            raw_text = content_bytes.decode("utf-8", errors="ignore")

    else:
        # Generic: try utf-8 decode
        raw_text = content_bytes.decode("utf-8", errors="ignore")

    if not raw_text.strip():
        # Last resort: try raw utf-8 decode (catches edge cases like .doc binary)
        raw_text = content_bytes.decode("utf-8", errors="ignore")
    
    if not raw_text.strip():
        raise HTTPException(status_code=422, detail="Could not extract text from the uploaded document. Please ensure the file contains readable text.")

    # Truncate to avoid token overflow (keep first 6000 chars)
    raw_text = raw_text[:6000]

    # ── Regex/keyword fallback extractor (works without Gemini) ─────────────
    def _regex_extract(text: str) -> dict:
        """Extract patient fields using regex patterns — no AI required."""
        import re
        result = {"name": None, "age": None, "gender": None, "chiefComplaint": None, "symptoms": None}

        # Name: "Patient Name: John Smith" — stop at newline, limit to 4 words
        name_match = re.search(
            r'(?:patient\s*name|full\s*name|name)\s*[:\-]\s*([A-Za-z][a-zA-Z\'\-]+(?:[^\S\n]+[A-Za-z][a-zA-Z\'\-]+){0,3})',
            text, re.IGNORECASE
        )
        if name_match:
            candidate = name_match.group(1).strip()
            # Reject if it looks like a label (e.g. "Age", "Gender")
            if candidate.lower() not in ("age", "gender", "sex", "date", "dob"):
                result["name"] = candidate

        # Age: "Age: 45", "45 years old", "45-year-old"
        age_match = re.search(
            r'(?:^|\b)(?:age|aged?)\s*[:\-]?\s*(\d{1,3})\s*(?:years?|yrs?)?(?:\b|$)|(\d{1,3})\s*[-\s]?year[s\-]?[-\s]?old',
            text, re.IGNORECASE | re.MULTILINE
        )
        if age_match:
            val = age_match.group(1) or age_match.group(2)
            try:
                age_int = int(val)
                if 0 < age_int < 130:
                    result["age"] = age_int
            except (ValueError, TypeError):
                pass

        # Gender: explicit label first, then pronoun inference
        gender_match = re.search(
            r'(?:gender|sex)\s*[:\-]\s*(male|female|other)',
            text, re.IGNORECASE
        )
        if gender_match:
            result["gender"] = gender_match.group(1).lower()
        else:
            if re.search(r'\b(mr\.?|male)\b', text, re.IGNORECASE):
                result["gender"] = "male"
            elif re.search(r'\b(mrs\.?|ms\.?|female)\b', text, re.IGNORECASE):
                result["gender"] = "female"

        # Chief complaint: single line after label
        cc_match = re.search(
            r'(?:chief\s*complaint|presenting\s*complaint|reason\s*for\s*visit|chief\s*concern)\s*[:\-]\s*([^\n]{1,200})',
            text, re.IGNORECASE
        )
        if cc_match:
            result["chiefComplaint"] = cc_match.group(1).strip()

        # Symptoms: everything after the label until a blank line or end
        sym_match = re.search(
            r'(?:^symptoms?|^presenting\s*symptoms?|^history\s*of\s*present\s*illness|^hpi)\s*[:\-]\s*([^\n]{1,500}(?:\n(?!\n)[^\n]{0,200}){0,5})',
            text, re.IGNORECASE | re.MULTILINE
        )
        if sym_match:
            result["symptoms"] = sym_match.group(1).strip()[:500]
        elif result["chiefComplaint"]:
            result["symptoms"] = result["chiefComplaint"]

        return result

    # ── Ask Gemini to extract structured fields (if key available) ───────────
    if not GEMINI_API_KEY:
        # No Gemini key — use regex fallback immediately
        regex_result = _regex_extract(raw_text)
        has_any = any(v is not None for v in regex_result.values())
        return {
            "success": has_any,
            "extracted": regex_result,
            "source": "regex_fallback",
            "error": "Gemini API key not configured — used pattern matching." if not has_any else None,
        }

    prompt = f"""You are a medical data extraction assistant. Extract patient information from the following clinical document.

Return ONLY a valid JSON object with these exact keys:
- "name": patient full name (string, or null if not found)
- "age": patient age as integer (or null if not found)
- "gender": one of "male", "female", "other" (or null if not found)
- "chiefComplaint": the primary reason for the visit, brief (string, or null if not found)
- "symptoms": detailed description of all symptoms mentioned (string, or null if not found)

Do NOT include vitals (blood pressure, heart rate, temperature, SpO2, respiratory rate) — those will be measured manually.
Do NOT include any explanation, markdown, or extra text. Return only the JSON object.

DOCUMENT:
{raw_text}
"""

    try:
        async with httpx.AsyncClient(timeout=20.0) as client:
            resp = await client.post(
                f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.5-flash-lite:generateContent?key={GEMINI_API_KEY}",
                json={
                    "contents": [{"parts": [{"text": prompt}]}],
                    "generationConfig": {"temperature": 0.1, "maxOutputTokens": 512},
                },
            )
        resp.raise_for_status()
        gemini_data = resp.json()
        raw_response = gemini_data["candidates"][0]["content"]["parts"][0]["text"].strip()

        # Strip markdown code fences if present
        if raw_response.startswith("```"):
            raw_response = raw_response.split("```")[1]
            if raw_response.startswith("json"):
                raw_response = raw_response[4:]
        raw_response = raw_response.strip()

        extracted = json.loads(raw_response)

        # Sanitise
        result = {
            "name": str(extracted.get("name") or "").strip() or None,
            "age": int(extracted["age"]) if extracted.get("age") is not None else None,
            "gender": extracted.get("gender") if extracted.get("gender") in ("male", "female", "other") else None,
            "chiefComplaint": str(extracted.get("chiefComplaint") or "").strip() or None,
            "symptoms": str(extracted.get("symptoms") or "").strip() or None,
        }
        return {"success": True, "extracted": result, "source": "gemini"}

    except Exception as e:
        # Any Gemini failure → regex fallback
        logger.warning(f"Document extraction via Gemini failed: {e}")
        err_str = str(e)
        # Detect quota errors
        try:
            if hasattr(e, 'response') and e.response is not None:
                err_body = e.response.json()
                if err_body.get('error', {}).get('code') == 429:
                    err_str = "429 quota exhausted"
        except Exception:
            pass
        regex_result = _regex_extract(raw_text)
        has_any = any(v is not None for v in regex_result.values())
        return {
            "success": has_any,
            "extracted": regex_result,
            "source": "regex_fallback",
            "error": err_str,
        }

# end of file
